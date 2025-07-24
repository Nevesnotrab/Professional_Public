import numpy as np
import random
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""
Creates a word search in a MS Word document page
"""

rows = 20
name = "Hanukkah"
maxit = 1000
it = 0

word = np.array(["HANUKKAH", "DREIDL", "LATKES", "MENORAH", "OIL", "FAMILY", "EIGHT", "MUSIC", "GRATITUDE", "CANDLES"])
chart = np.empty((rows, rows), dtype = str)
check = np.empty((rows, rows), dtype = bool)
for i in range(rows):
    for j in range(rows):
        check[i,j] = False

maxsize = 0

#Finds the length of the longest word
for i in range(len(word)):
    if(len(word[i]) > maxsize):
        maxsize = len(word[i])

#Returns a random letter
def RandomLetter():
    letter = random.randrange(1, 26)
    return chr(letter + 64)

#Checks to see if the word can go there
def IsFilled(row, col, dire, leng):
    for i in range(0, leng):
        if(dire == 0):
            if(check[row,col] == True):
                return True
            else:
                row -= 1
                if (row < 0):
                    return True
        elif(dire == 1):
            if(check[row, col] == True):
                return True
            else:
                row -= 1
                col += 1
                if (row < 0) or (col > rows - 1):
                    return True
        elif(dire == 2):
            if(check[row, col] == True):
                return True
            else:
                col += 1
                if (col > rows - 1):
                    return True
        elif(dire == 3):
            if(check[row, col] == True):
                return True
            else:
                row += 1
                col += 1
                if (row > rows - 1) or (col > rows - 1):
                    return True
        elif(dire == 4):
            if(check[row, col] == True):
                return True
            else:
                row += 1
                if(row > rows - 1):
                    return True
        elif(dire == 5):
            if(check[row, col] == True):
                return True
            else:
                row += 1
                col -= 1
                if(row > rows - 1) or (col < 0):
                    return True
        elif(dire == 6):
            if(check[row, col] == True):
                return True
            else:
                col -= 1
                if(col < 0):
                    return True
        else:
            if(check[row, col] == True):
                return True
            else:
                row -= 1
                col -= 1
                if(row < 0) or (col < 0):
                    return True
    return False

#Takes a word, finds a spot for it, and fills it in
def InsertWord(word):
    global it
    it += 1
    if(it > maxit):
        print("Exceeded maximum iterations")
        return
    leng = len(word)
    row = random.randrange(0, rows - 1)
    col = random.randrange(0, rows - 1)
    dire = random.randrange(0, 7)
    
    if(IsFilled(row, col, dire, leng) == True):
        InsertWord(word)
    else:
        for i in range(0, leng):
            chart[row, col] = word[i]
            check[row, col] = True
            if(dire == 0):
                row -= 1
            elif(dire == 1):
                row -= 1
                col += 1
            elif(dire == 2):
                col += 1
            elif(dire == 3):
                row += 1
                col += 1
            elif(dire == 4):
                row += 1
            elif(dire == 5):
                row += 1
                col -= 1
            elif(dire == 6):
                col -= 1
            else:
                row -= 1
                col -= 1


#Assigns a random letter to each spot in the chart
for i in range(rows):
    for j in range(rows):
        chart[i,j] = RandomLetter()
        
for i in range(0, len(word)):
    InsertWord(word[i])
    
document = Document()
document.add_heading(name, 0)

#Prints the chart
for i in range(rows):
    for j in range(rows):
        print(chart[i, j], end = '', flush = True)
        if(((j + 1) % rows) != 0):
            print('   ', end = '', flush = True)
    print('\n')

table = document.add_table(rows = 1, cols = rows)

for i in range(0, rows):
    row_cells = table.add_row().cells
    for j in range(0, rows):
        row_cells[j].text = chart[i,j]
        
listwords = '\n\n'
for i in range(len(word)):
    listwords += word[i]
    if (i != len(word) - 1):
        listwords += '        '

paragraph = document.add_paragraph(listwords)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.save(name + '.docx')